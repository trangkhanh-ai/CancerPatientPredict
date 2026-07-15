# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Pt


DOCS = [
    Path("final+spark/BaoCao_Final_Spark_MLlib.docx"),
    Path("tai_lieu_bao_cao/BaoCao_Final_Spark_MLlib.docx"),
]


P145 = (
    "Kiến trúc hệ thống được thiết kế theo hướng pipeline dữ liệu. Dữ liệu ban đầu được chuẩn hóa "
    "từ Excel sang CSV để Hadoop Streaming đọc thuận tiện. File CSV được upload lên HDFS, sau đó "
    "MapReduce job thống kê phân bố nhãn, nhóm tuổi và mức hút thuốc. Kết quả thống kê được import "
    "vào MongoDB để dashboard truy vấn. Song song, dữ liệu sạch được dùng để huấn luyện mô hình bằng "
    "Apache Spark MLlib. PipelineModel của Logistic Regression chính thức được lưu ra thư mục "
    "model_spark/ và được FastAPI nạp khi khởi động."
)
P161 = (
    "Chia train/test theo group-aware split dựa trên feature_signature, cố định seed=42 và kiểm tra "
    "overlap=0 để tránh rò rỉ dữ liệu."
)
P162 = (
    "Huấn luyện Logistic Regression đa lớp trong Pipeline Spark MLlib (VectorAssembler, "
    "StandardScaler, LogisticRegression) và đánh giá bằng Accuracy, Precision, Recall, F1-score."
)
P192 = (
    "Giai đoạn Spark: sau khi MapReduce tổng hợp dữ liệu, Spark đảm nhận hai nhánh song song. "
    "Spark SQL đọc dữ liệu (từ MongoDB/HDFS hoặc CSV) và truy vấn thống kê bằng GROUP BY theo nhãn, "
    "nhóm tuổi, trung bình các yếu tố nguy cơ. Spark MLlib huấn luyện Logistic Regression chính thức "
    "trên tập group-aware, kiểm tra feature_signature không chồng lấn giữa train và test, rồi lưu "
    "PipelineModel để FastAPI phục vụ trực tuyến. Nhờ vậy toàn bộ khâu truy vấn, huấn luyện và phục vụ "
    "mô hình đều chạy trên cùng một nền tảng phân tán."
)
P220 = (
    "Kết quả thực nghiệm trong Bảng 14 được thu trực tiếp từ quá trình huấn luyện thật bằng Apache "
    "Spark MLlib trên tập dữ liệu đã tiền xử lý (1.000 bệnh nhân). Do phát hiện 848 dòng trùng vector "
    "đặc trưng (chỉ 152 feature_signature duy nhất), báo cáo không dùng randomSplit theo dòng làm kết "
    "quả chính thức. Dữ liệu được chia theo group-aware split với group = feature_signature, seed = 42, "
    "train = 697 dòng, test = 303 dòng và signature_overlap = 0. Mô hình chính thức là Logistic "
    "Regression đa lớp trong Pipeline Spark MLlib (VectorAssembler, StandardScaler, LogisticRegression). "
    "Trên tập test group-aware, mô hình đạt Accuracy, Precision, Recall và F1 đều bằng 1,000. Ma trận "
    "nhầm lẫn là ma trận chéo hoàn hảo với Low = 87, Medium = 90 và High = 126. Kết quả tuyệt đối phản "
    "ánh đặc tính của bộ dữ liệu học thuật: các nhãn Level gần như được xác định bởi tổ hợp 23 đặc trưng "
    "và không có feature_signature nào gắn nhiều nhãn. Vì vậy kết quả này chứng minh pipeline đánh giá "
    "hợp lệ trên dữ liệu hiện có, không được diễn giải như bằng chứng về năng lực tổng quát hóa trên dữ "
    "liệu lâm sàng thực tế."
)
P224 = (
    "Ma trận nhầm lẫn của Logistic Regression trên tập kiểm thử 303 mẫu là ma trận chéo hoàn hảo: "
    "toàn bộ bản ghi nằm trên đường chéo chính (Low 87, Medium 90, High 126), không có dự đoán sai. "
    "Điểm quan trọng là các dòng có cùng feature_signature luôn nằm cùng một phía train hoặc test, nên "
    "kết quả không phụ thuộc vào rò rỉ do bản ghi trùng vector xuất hiện đồng thời ở hai tập."
)
P225 = (
    "Các đặc trưng như coughing_of_blood, obesity, wheezing, passive_smoker và fatigue vẫn có ý nghĩa "
    "diễn giải vì trung bình của chúng tăng rõ giữa các mức nguy cơ trong dataset. Tuy nhiên, bảng này "
    "chỉ dùng để hỗ trợ giải thích dữ liệu và không thay thế bằng chứng y khoa độc lập; khi trình bày "
    "trước hội đồng cần nhấn mạnh đây là kết quả học từ bộ dữ liệu thực nghiệm."
)
P238 = (
    "Về mô hình, kết quả chính thức dùng group-aware split theo feature_signature và Logistic Regression "
    "Spark MLlib như sau:"
)
P240 = (
    "Kết quả chạy lại bằng train_spark.py sau chỉnh sửa cho thấy rows_train = 697, rows_test = 303, "
    "signature_overlap = 0, Accuracy = 1,000 và F1 = 1,000. Việc bỏ randomSplit khỏi đường huấn luyện "
    "chính thức giúp báo cáo, file ml_results_spark.json và PipelineModel nói cùng một câu chuyện: "
    "Logistic Regression là mô hình phục vụ, còn group-aware split là phương pháp đánh giá chống leakage."
)
P254 = (
    "Đã trình bày quy trình huấn luyện Logistic Regression bằng Spark MLlib trên split group-aware và "
    "lưu PipelineModel cho API."
)
APPENDIX_INTRO = (
    "Script huấn luyện thật bằng Apache Spark MLlib trên cancer_patients_ml_ready.csv: đọc dữ liệu bằng "
    "Spark DataFrame, tạo feature_signature, chia group-aware overlap=0, huấn luyện Logistic Regression "
    "và lưu PipelineModel cho FastAPI."
)
APPENDIX = [
    "# -*- coding: utf-8 -*-",
    '"""train_spark.py - bản chính thức sau chỉnh sửa."""',
    "Các điểm chính:",
    "- Đọc cancer_patients_ml_ready.csv bằng Spark DataFrame.",
    "- Tạo feature_signature = SHA-256 của 23 đặc trưng theo đúng thứ tự.",
    "- Chia group-aware theo feature_signature, seed=42, train_ratio=0,70.",
    "- Kiểm tra signature_overlap = 0 trước khi huấn luyện.",
    "- Huấn luyện Pipeline: VectorAssembler -> StandardScaler -> LogisticRegression(multinomial, regParam=0.0).",
    "- Lưu PipelineModel vào model_spark/ và kết quả vào ml_results_spark.json.",
    "",
    "Kết quả chạy thử Spark:",
    "[BASELINE] rows=1000 unique_signatures=152 duplicated_feature_rows=848 conflicts=0",
    "[GROUP SPLIT] rows_train=697 rows_test=303 signature_overlap=0",
    "[TEST SUPPORT] {'Low': 87, 'Medium': 90, 'High': 126}",
    "[LR] accuracy=1.0000 f1=1.0000 confusion=[[87, 0, 0], [0, 90, 0], [0, 0, 126]]",
    "",
    "Ghi chú khi chạy trên Windows:",
    "$env:JAVA_HOME='C:\\Program Files\\Eclipse Adoptium\\jdk-17.0.19.10-hotspot'",
    "python train_spark.py",
]


def set_paragraph_text(paragraph, text, font="Times New Roman", size=12):
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = font
        run.font.size = Pt(size)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_row(row):
    element = row._tr
    element.getparent().remove(element)


def shrink_table(table, keep_rows):
    while len(table.rows) > keep_rows:
        delete_row(table.rows[-1])


def set_cell(cell, text, bold=False):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = bold


def update_table_13(table):
    shrink_table(table, 2)
    header = ["Mô hình", "Accuracy", "Precision", "Recall", "F1", "Ghi chú"]
    row = ["Logistic Regression", "1,000", "1,000", "1,000", "1,000", "group-aware, test=303"]
    for idx, text in enumerate(header):
        set_cell(table.rows[0].cells[idx], text, bold=True)
    for idx, text in enumerate(row):
        set_cell(table.rows[1].cells[idx], text)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def update_confusion(table):
    values = [
        ["Thực tế \\ Dự đoán", "Low", "Medium", "High"],
        ["Low", "87", "0", "0"],
        ["Medium", "0", "90", "0"],
        ["High", "0", "0", "126"],
    ]
    for row_idx, row in enumerate(values):
        for col_idx, text in enumerate(row):
            set_cell(table.rows[row_idx].cells[col_idx], text, bold=(row_idx == 0 or col_idx == 0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def update_table_19(table):
    shrink_table(table, 2)
    header = ["Mô hình", "Accuracy", "Precision", "Recall", "F1"]
    row = ["Logistic Regression", "1,000", "1,000", "1,000", "1,000"]
    for idx, text in enumerate(header):
        set_cell(table.rows[0].cells[idx], text, bold=True)
    for idx, text in enumerate(row):
        set_cell(table.rows[1].cells[idx], text)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_by_index(doc, index, text):
    if index < len(doc.paragraphs):
        set_paragraph_text(doc.paragraphs[index], text)


def update_doc(path):
    doc = Document(path)
    replacements = {
        145: P145,
        161: P161,
        162: P162,
        192: P192,
        220: P220,
        222: "Bảng 15: Ma trận nhầm lẫn của Logistic Regression trên tập group-aware",
        223: "Bảng 16: Một số đặc trưng phân biệt chính trong dữ liệu",
        224: P224,
        225: P225,
        238: P238,
        239: "Bảng 20: Kết quả Spark MLlib chính thức.",
        240: P240,
        254: P254,
    }
    for index, text in replacements.items():
        set_by_index(doc, index, text)

    update_table_13(doc.tables[13])
    update_confusion(doc.tables[14])
    update_table_19(doc.tables[19])

    start = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("G. train_spark.py"):
            start = idx
            break
    if start is None:
        raise RuntimeError(f"Cannot find appendix G heading in {path}")

    set_paragraph_text(doc.paragraphs[start + 1], APPENDIX_INTRO)
    keep_until = start + 2
    while len(doc.paragraphs) > keep_until:
        delete_paragraph(doc.paragraphs[-1])
    for line in APPENDIX:
        if line.startswith(("[", "$", "python", "-", "#", '"')):
            set_paragraph_text(doc.add_paragraph(), line, font="Consolas", size=10)
        else:
            set_paragraph_text(doc.add_paragraph(), line)

    doc.save(path)


for doc_path in DOCS:
    update_doc(doc_path)
    print(f"updated {doc_path}")
