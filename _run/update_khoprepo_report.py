# -*- coding: utf-8 -*-
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Pt


SRC = Path(r"C:\Users\dzyuu\Downloads\BaoCao_UngThu_BigData_KhopRepo.docx")
OUT_DOWNLOADS = Path(r"C:\Users\dzyuu\Downloads\BaoCao_UngThu_BigData_KhopRepo_DaBoSung.docx")
OUT_REPO = Path(r"D:\1BigDataproject1\final\tai_lieu_bao_cao\BaoCao_UngThu_BigData_KhopRepo_DaBoSung.docx")
AUDIT = Path(r"D:\1BigDataproject1\final\BAO_CAO_KHOP_REPO_GIAI_THICH_CHINH_SUA.md")


REPLACEMENTS = {
    23: (
        "Kho mã nguồn trong workspace D:\\1BigDataproject1\\final đã hiện thực đủ năm module: "
        "làm sạch và chuẩn hoá dữ liệu, hai Hadoop MapReduce job, MongoDB, pipeline PySpark ML, "
        "FastAPI và ứng dụng C# WinForms năm tab. Báo cáo dùng mã nguồn, artifact và tài liệu vận hành "
        "trong repo làm căn cứ; trạng thái ‘có source’ được tách khỏi trạng thái ‘đã chạy lại trong môi trường rà soát’."
    ),
    24: (
        "Artifact canonical trong repo cho thấy 152 feature_signature duy nhất trên 1.000 hồ sơ, tức 848 hồ sơ "
        "lặp vector 23 đặc trưng và không có xung đột nhãn. Phép chia group-aware xác định bằng SHA-256 với "
        "seed 42 tạo train 697 dòng/106 nhóm và test 303 dòng/46 nhóm, overlap bằng 0. Repo đã được chạy lại "
        "bằng Spark MLlib sau khi đồng bộ split canonical: Logistic Regression và Random Forest đều đạt "
        "Accuracy = F1 = 1,0000; PipelineModel phục vụ hiện hành trong models/current đã được cập nhật sang "
        "Logistic Regression group-aware."
    ),
    117: (
        "Phạm vi đối chiếu là workspace hiện tại D:\\1BigDataproject1\\final. Repo chứa dữ liệu gốc/đã xử lý, "
        "mã nguồn cho cả năm module, model artifact hiện hành, metrics canonical, tài liệu kiến trúc, API và runbook. "
        "Báo cáo chỉ gọi một thành phần là ‘đã kiểm chứng’ khi repo có artifact hoặc trạng thái chạy tương ứng; "
        "MongoDB và bản build WinForms vẫn cần chạy lại với dịch vụ/SDK bên ngoài trước buổi demo."
    ),
    122: (
        "Mã nguồn đã có /predict, /patients, training PySpark, PipelineModel và đầy đủ solution WinForms. "
        "Spark canonical đã chạy lại thành công trên JDK 17; giới hạn còn lại là chưa chạy lại end-to-end "
        "MongoDB + FastAPI + WinForms trong lần rà soát tài liệu này."
    ),
    174: (
        "Kiến trúc được đối chiếu gồm năm module liên kết bằng schema snake_case dùng chung. Module 1 làm sạch, "
        "kiểm định và ghi HDFS/CSV; Module 2 chạy hai MapReduce job phân bố và tương quan; MongoDB quản lý sáu "
        "collection; Module 3 huấn luyện và lưu PipelineModel Logistic Regression; Module 4 cung cấp 11 route "
        "dưới /api/v1; Module 5 là WinForms .NET 8 với năm tab. Hình 2 mô tả các thành phần thực có trong repo, "
        "không dùng nét đứt để ám chỉ thiếu source."
    ),
    211: (
        "Repo được đối chiếu yêu cầu Python 3.10+, PySpark 4.x với JDK 17+, MongoDB cục bộ và .NET 8. "
        "Dữ liệu gốc data/raw/cancer_patient_data_sets.xlsx có 1.000 × 25; data/processed/"
        "cancer_patients_ml_ready.csv có 1.000 × 26 vì bổ sung level_encoded. Artifact distributions.tsv, "
        "correlation.tsv, metrics.json, confusion_group_aware.csv, models/current/metadata.json và "
        "final+spark/ml_results_spark.json đều đã có trong workspace."
    ),
    228: (
        "Repo hiện đã có một bộ số chính thức thống nhất. Artifact canonical do ml_analysis.py tạo trên holdout "
        "group-aware 697/303, overlap 0, ghi Accuracy = Macro-F1 = 1,0000 cho Logistic Regression kết hợp "
        "StandardScaler. Sau rà soát, Spark MLlib cũng đã chạy lại cùng split canonical: rows_train = 697, "
        "rows_test = 303, Accuracy = F1 = 1,0000 và confusion [[87,0,0],[0,90,0],[0,0,126]]. "
        "models/current/metadata.json và final+spark/ml_results_spark.json cùng trình bày Logistic Regression "
        "là mô hình phục vụ."
    ),
    231: (
        "Kết quả canonical 1,0000 vẫn không chứng minh khả năng tổng quát hoá lâm sàng. Dù overlap nhóm bằng 0, "
        "dữ liệu chỉ có 152 signature và không có signature nào mang nhãn xung đột; nhãn gần như tách biệt theo "
        "các chỉ số. Vì vậy điểm tuyệt đối phản ánh đặc tính separable của dataset học thuật, không phải cam kết "
        "hiệu năng trên dữ liệu bệnh viện thực tế."
    ),
    249: (
        "Ba ảnh là minh hoạ giao diện do nhóm cung cấp. Repo hiện có đủ năm tab và luồng gọi API tương ứng, "
        "nhưng khi chụp lại bản cuối nên hiển thị model_run_id = lr-group-aware-20260715, dataset_version = "
        "v20260715 và latency từ response để ảnh khớp chính xác với artifact đang phục vụ."
    ),
    253: (
        "Đối chiếu workspace hiện tại cho thấy repo đã có cấu trúc end-to-end: dữ liệu và kiểm định, hai MapReduce "
        "job, sáu collection MongoDB, pipeline Logistic Regression, 11 route FastAPI và ứng dụng WinForms năm tab. "
        "Các artifact dữ liệu/metrics và Spark model hiện có đủ để kiểm tra định lượng; các dịch vụ phụ thuộc "
        "MongoDB, FastAPI và .NET vẫn cần chạy lại trước buổi demo."
    ),
    262: (
        "PySpark đã được huấn luyện lại sau khi đồng bộ hàm chia canonical; metrics canonical 1,0000 và metadata "
        "PipelineModel hiện hành đều cùng một câu chuyện: group-aware split 697/303, overlap 0, Logistic Regression."
    ),
    265: (
        "MapReduce đã có artifact dev, nhưng chưa có benchmark cụm; MongoDB, FastAPI và WinForms vẫn cần chạy lại "
        "end-to-end trong môi trường demo."
    ),
    269: "Đánh giá repeated group-aware holdout trên nhiều seed và tập dữ liệu ngoài.",
    270: "Tối ưu MongoDB aggregation; benchmark HDFS/MapReduce/Spark và chụp lại giao diện từ bản build cuối.",
    290: (
        "[19] Repo cục bộ D:\\1BigDataproject1\\final, kiểm tra ngày 15/07/2026: module source, metrics, "
        "model artifact và runbook của đồ án Cancer Big Data."
    ),
    305: (
        "Module 3: split.py, train.py, PipelineModel hiện hành và metrics đã đồng bộ; Spark canonical đã chạy "
        "lại với train=697/test=303, overlap=0, Accuracy=1,0000."
    ),
    306: (
        "MongoDB: script tạo schema validator, index và bulk upsert idempotent cho sáu collection đã có; verify "
        "cần một instance MongoDB đang chạy."
    ),
    307: (
        "Module 4: FastAPI có 11 route dưới /api/v1, validation Pydantic, whitelist query, CORS cấu hình và "
        "model service nạp PipelineModel một lần."
    ),
    308: (
        "Module 5: solution WinForms .NET 8 có MainForm, năm control, shared HttpClient và health polling 30 giây; "
        "cần build/run và chụp lại ảnh từ phiên bản cuối."
    ),
}


def style_paragraph(paragraph, size=12):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def set_paragraph(doc, idx, text):
    if idx >= len(doc.paragraphs):
        raise IndexError(f"paragraph index out of range: {idx}")
    doc.paragraphs[idx].text = text
    style_paragraph(doc.paragraphs[idx])


def set_cell(cell, text, bold=False, size=10):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold


def ensure_rows(table, rows):
    while len(table.rows) < rows:
        table.add_row()
    while len(table.rows) > rows:
        element = table.rows[-1]._tr
        element.getparent().remove(element)


def update_table(table, data):
    ensure_rows(table, len(data))
    for row_idx, row in enumerate(data):
        for col_idx, text in enumerate(row):
            set_cell(table.rows[row_idx].cells[col_idx], text, bold=(row_idx == 0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def update_report():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    OUT_REPO.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT_DOWNLOADS)
    shutil.copy2(SRC, OUT_REPO)

    for out_path in (OUT_DOWNLOADS, OUT_REPO):
        doc = Document(out_path)
        for idx, text in REPLACEMENTS.items():
            set_paragraph(doc, idx, text)

        update_table(
            doc.tables[19],
            [
                ["Artifact", "Tập đánh giá", "Group overlap", "Accuracy", "F1", "Trạng thái"],
                ["Canonical metrics.json", "303 dòng / 46 nhóm", "0", "1,0000", "Macro-F1 1,0000", "scikit-learn proxy trên split canonical"],
                ["Spark MLlib models/current", "303 dòng", "0", "1,0000", "Weighted-F1 1,0000", "đã chạy lại bằng PySpark/JDK 17"],
                ["Spark MLlib final+spark", "303 dòng", "0", "1,0000", "F1 1,0000", "ml_results_spark.json chính thức"],
            ],
        )
        update_table(
            doc.tables[21],
            [
                ["Yêu cầu", "Bằng chứng trong repo", "Trạng thái", "Kết luận"],
                ["Giao diện thân thiện", "MainForm + 5 control WinForms .NET 8", "Đủ source", "Cần build/run và chụp lại bản cuối"],
                ["Lưu trữ MongoDB", "Validator, index, bulk upsert; 6 collection", "Đủ source", "Cần verify với MongoDB đang chạy"],
                ["Dự đoán mức độ", "train.py, PipelineModel, POST /predict", "Có source + artifact đã đồng bộ", "Logistic Regression group-aware là bản chính thức"],
                ["Thống kê bệnh nhân", "MapReduce artifacts, /stats, dashboard", "Có source + artifact", "Đúng phạm vi chức năng"],
                ["Tìm kiếm/lọc", "/patients, export, detail; PatientSearchControl", "Đủ source", "Có pagination và whitelist filter"],
                ["Chất lượng dữ liệu", "compute_quality, /quality, DataQualityControl", "Có source + metrics", "Phân biệt row validity và completeness"],
            ],
        )
        update_table(
            doc.tables[22],
            [
                ["Vấn đề còn lại", "Bằng chứng", "Ảnh hưởng", "Cách khắc phục"],
                ["Dịch vụ ngoài chưa chạy lại end-to-end", "MongoDB, FastAPI, WinForms", "Chưa xác nhận demo trên máy nộp", "Chạy runbook, lưu log và ảnh màn hình"],
                ["Quy mô nhỏ", "1.000 dòng; 152 signature", "Không chứng minh hiệu năng Big Data", "Benchmark tập lớn trên cụm"],
                ["Giới hạn lâm sàng", "Dữ liệu công khai, 0 xung đột nhãn", "Không suy rộng thành chẩn đoán", "Giữ disclaimer; đánh giá dữ liệu ngoài"],
                ["Model version cần xuất hiện trong demo", "metadata có model_run_id/dataset_version", "Ảnh UI cũ dễ lệch artifact", "Chụp lại UI với lr-group-aware-20260715"],
                ["Mongo aggregation còn tải về Python", "data_services.get_stats/get_quality", "Chưa tối ưu khi dữ liệu lớn", "Đẩy aggregation xuống MongoDB khi mở rộng"],
                ["Bảo mật mức đồ án", "CORS/config có nhưng chưa auth", "Không dùng cho môi trường thật", "Thêm auth, rate limit, audit log"],
            ],
        )
        update_table(
            doc.tables[23],
            [
                ["Hạng mục", "Trạng thái theo repo", "Việc cần làm trước khi nộp"],
                ["Báo cáo DOCX", "Đã đồng bộ với repo hiện tại", "Cập nhật field trong Word và đọc soát lần cuối"],
                ["Báo cáo PDF", "Cần xuất từ DOCX đã đồng bộ", "Kiểm tra trang, hình, bảng và font"],
                ["Repo/local workspace", "D:\\1BigDataproject1\\final", "Đóng gói đúng thư mục đã chạy"],
                ["Dataset gốc", "Có data/raw/*.xlsx", "Giữ nguồn và giấy phép/URL"],
                ["Dataset processed", "Có ML-ready CSV 1.000 × 26", "Kiểm tra checksum khi đóng gói"],
                ["MapReduce", "2 job + 2 TSV artifact", "Lưu lệnh chạy và log demo"],
                ["PySpark/model", "Đã rerun canonical; models/current đã đồng bộ", "Giữ metadata.json và ml_results_spark.json"],
                ["MongoDB", "Có validator/index/import/verify", "Chạy verify trên instance thực"],
                ["FastAPI", "Có 11 route /api/v1", "Chạy smoke/integration test"],
                ["WinForms", "Có solution .NET 8 + 5 tab", "Build/run và chụp lại UI"],
                ["Tài liệu chạy", "README, ARCHITECTURE, API, RUNBOOK đã sửa 11 endpoint", "Đọc lại trước nộp"],
                ["Bảo mật/benchmark", "Mức đồ án học thuật", "Thêm auth/rate limit và benchmark nếu mở rộng"],
            ],
        )
        doc.save(out_path)


def write_audit():
    AUDIT.write_text(
        """# Giải thích chỉnh sửa báo cáo khớp repo

## File đầu vào

- `C:\\Users\\dzyuu\\Downloads\\BaoCao_UngThu_BigData_KhopRepo.docx`

## File đã tạo

- `C:\\Users\\dzyuu\\Downloads\\BaoCao_UngThu_BigData_KhopRepo_DaBoSung.docx`
- `tai_lieu_bao_cao/BaoCao_UngThu_BigData_KhopRepo_DaBoSung.docx`

## Những điểm đã phát hiện

1. Báo cáo cũ ghi `commit e94b59e8fa`; workspace hiện tại không phải git repo, nên đổi sang mô tả repo cục bộ `D:\\1BigDataproject1\\final`.
2. Báo cáo cũ nói PySpark canonical chưa chạy lại và `models/current` còn 0,9668/test=301.
3. Repo hiện tại đã được chạy lại Spark canonical:
   - `rows_train=697`, `rows_test=303`
   - `signature_overlap=0`
   - Logistic Regression: `accuracy=1.0000`, `f1=1.0000`
   - confusion matrix: `[[87,0,0],[0,90,0],[0,0,126]]`
4. `models/current/metadata.json`, `models/runs/logistic_regression/metrics.json`, `_run/ml_results_spark.json` và `final+spark/ml_results_spark.json` đã cùng nói Logistic Regression group-aware.
5. API code có 11 route dưới `/api/v1`, nên sửa mô tả 10 endpoint thành 11 endpoint trong tài liệu repo và báo cáo.

## Những gì đã sửa trong DOCX

- Cập nhật phần tóm tắt, phạm vi, kiến trúc, kết quả mô hình, kết luận và phụ lục trạng thái.
- Cập nhật bảng Artifact mô hình: canonical metrics, Spark `models/current`, Spark `final+spark` đều 303 test, overlap 0, acc 1.0.
- Cập nhật bảng đối chiếu yêu cầu đề bài: dự đoán mức độ đã có artifact đồng bộ.
- Cập nhật bảng vấn đề còn lại: bỏ “chưa rerun PySpark canonical”, thay bằng các việc thật còn phải làm trước demo như chạy MongoDB/FastAPI/WinForms và chụp lại UI.
- Cập nhật checklist nộp: PySpark/model đã rerun canonical; cần giữ metadata và JSON kết quả.

## Lệnh Spark đã chạy

```powershell
$env:JAVA_HOME='C:\\Program Files\\Eclipse Adoptium\\jdk-17.0.19.10-hotspot'
$env:PYTHONUTF8='1'
$env:PYTHONPATH='src'
python src\\ml\\split.py --input data\\processed\\cancer_patients_ml_ready.csv --out data\\processed
python src\\ml\\train.py --split data\\processed\\split_manifest.parquet --out models
```
""",
        encoding="utf-8",
    )


if __name__ == "__main__":
    update_report()
    write_audit()
    print(OUT_DOWNLOADS)
    print(OUT_REPO)
    print(AUDIT)
